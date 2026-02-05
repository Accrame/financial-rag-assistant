# Document loading for financial PDFs, text files, etc.

import hashlib
from pathlib import Path

import pdfplumber
from langchain_core.documents import Document
from pypdf import PdfReader


class FinancialDocumentLoader:
    """Loads financial docs (PDFs, text, docx) with metadata."""

    def __init__(self, extract_tables=True):
        self.extract_tables = extract_tables

    def load(self, file_path):
        """Load a document and return list of Document objects."""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self._load_pdf(path)
        elif suffix == ".txt":
            return self._load_text(path)
        elif suffix in [".doc", ".docx"]:
            return self._load_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _load_pdf(self, path, doc_id=None):
        """Load PDF with page-level granularity."""
        if doc_id is None:
            doc_id = self._generate_doc_id(path)

        # use pdfplumber for tables, fall back to pypdf otherwise
        if self.extract_tables:
            return self._load_pdf_with_pdfplumber(path, doc_id)
        else:
            return self._load_pdf_with_pypdf(path, doc_id)

    def _load_pdf_with_pypdf(self, path, doc_id):
        reader = PdfReader(str(path))
        documents = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""

            if text.strip():
                doc = Document(
                    page_content=text,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "doc_id": doc_id,
                        "page": page_num,
                        "total_pages": len(reader.pages),
                        "file_type": "pdf",
                    },
                )
                documents.append(doc)

        return documents

    def _load_pdf_with_pdfplumber(self, path, doc_id):
        """Load PDF using pdfplumber (better for tables)."""
        documents = []

        with pdfplumber.open(str(path)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract main text
                text = page.extract_text() or ""

                # Extract tables if present
                tables = page.extract_tables()
                table_text = ""

                if tables:
                    for i, table in enumerate(tables):
                        table_text += f"\n[Table {i + 1}]\n"
                        for row in table:
                            # Clean None values
                            cleaned_row = [str(cell) if cell else "" for cell in row]
                            table_text += " | ".join(cleaned_row) + "\n"

                combined_text = text
                if table_text:
                    combined_text += "\n\n--- Tables ---" + table_text

                if combined_text.strip():
                    doc = Document(
                        page_content=combined_text,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "doc_id": doc_id,
                            "page": page_num,
                            "total_pages": len(pdf.pages),
                            "file_type": "pdf",
                            "has_tables": bool(tables),
                        },
                    )
                    documents.append(doc)

        return documents

    def _load_text(self, path):
        """Load a plain text file."""
        doc_id = self._generate_doc_id(path)

        with open(path, encoding="utf-8") as f:
            text = f.read()

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "doc_id": doc_id,
                    "file_type": "txt",
                },
            )
        ]

    def _load_docx(self, path):
        from docx import Document as DocxDocument

        doc_id = self._generate_doc_id(path)
        docx = DocxDocument(str(path))

        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs)

        return [
            Document(
                page_content=text,
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "doc_id": doc_id,
                    "file_type": "docx",
                },
            )
        ]

    def _generate_doc_id(self, path):
        """Hash-based doc ID for deduplication."""
        with open(path, "rb") as f:
            content = f.read()
        return hashlib.md5(content).hexdigest()[:12]

    def load_directory(self, directory, extensions=None):
        """Load all supported docs from a directory."""
        if extensions is None:
            extensions = [".pdf", ".txt", ".docx"]

        path = Path(directory)
        documents = []

        for ext in extensions:
            for file_path in path.glob(f"*{ext}"):
                try:
                    docs = self.load(str(file_path))
                    documents.extend(docs)
                except Exception as e:
                    # TODO: add proper logging instead of print
                    print(f"Error loading {file_path}: {e}")

        # print(f"DEBUG: loaded {len(documents)} docs from {directory}")
        return documents


if __name__ == "__main__":
    # Quick test
    loader = FinancialDocumentLoader()
    print("Document loader initialized successfully")
    print("Supported formats: PDF, TXT, DOCX")
